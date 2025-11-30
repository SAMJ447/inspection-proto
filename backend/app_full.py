import os
import json
import uuid
import base64
import re
from io import BytesIO
from typing import List, Optional, Dict, Any

from fastapi import (
    FastAPI,
    UploadFile,
    File,
    HTTPException,
    Query,
    Body,
    Form,
)
from fastapi.responses import HTMLResponse, StreamingResponse
from starlette.responses import JSONResponse
from pydantic import BaseModel

from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.utils import ImageReader

from openai import OpenAI
from docx import Document
from docx.shared import Inches

# ============================================================
# FastAPI app
# ============================================================

app = FastAPI()

# ============================================================
# Paths / dirs
# ============================================================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
ANNOTATION_DIR = os.path.join(BASE_DIR, "annotations")
TEMPLATE_DIR = os.path.join(BASE_DIR, "templates")
COMPANY_TEMPLATE_DIR = os.path.join(TEMPLATE_DIR, "companies")
GLOBAL_TEMPLATE_DIR = os.path.join(TEMPLATE_DIR, "global")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(ANNOTATION_DIR, exist_ok=True)
os.makedirs(TEMPLATE_DIR, exist_ok=True)
os.makedirs(COMPANY_TEMPLATE_DIR, exist_ok=True)
os.makedirs(GLOBAL_TEMPLATE_DIR, exist_ok=True)

DEFAULT_TEMPLATE_PATH = os.path.join(TEMPLATE_DIR, "default_report_template.docx")

# ============================================================
# OpenAI client
# ============================================================

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    print("WARNING: OPENAI_API_KEY is not set. /generate-report and /generate-report-from-images will fail.")

client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None

# ============================================================
# Trade style prompts
# ============================================================

TRADE_STYLE_PROMPTS: Dict[str, str] = {
    "welding": """
Write in the tone of a NYC special inspection welding report.

- Mention AWS D1.1 or project specifications where reasonable.
- Reference member sizes and weld sizes explicitly when visible.
- Confirm continuity of fillet welds, absence of visible cracks, porosity, undercut, or slag.
- Clearly state ACCEPTED vs REJECTED conditions in professional language.
""",
    "bolting": """
Write in the tone of a NYC special inspection bolting (HSB) report.

- Mention A325 / A490 or project bolt specs when appropriate.
- Refer to installation method (snug-tight, fully pretensioned), and any visual verification of bolt heads, washers, and plies.
- Call out any missing bolts, loose bolts, or discrepancies clearly.
""",
    "detail": """
Write as a detail / connection verification special inspection.

- Focus on verifying that the built condition matches the referenced detail(s).
- Reference detail numbers (e.g., 5/S3.1) and any deviations from the design.
""",
    "cold_formed_steel": """
Write as a cold-formed steel framing special inspection.

- Focus on stud sizes, spacing, connections, bracing, and anchors.
- Reference applicable standards and project details.
""",
    "sprinkler": """
Write as a fire sprinkler special inspection.

- Reference NFPA standards where appropriate.
- Focus on pipe sizes, hanger spacing, valve locations, head types, and coverage.
""",
    "standpipe": """
Write as a standpipe special inspection.

- Address riser piping, valves, hose connections, and required signage.
""",
    "hvac": """
Write as an HVAC special inspection.

- Focus on duct routing, fire/smoke dampers, equipment supports, and penetration details as applicable.
""",
    "fire_resistance_penetration": """
Write as a fire-resistance rated penetration / firestopping special inspection.

- Focus on penetration type, assembly rating, tested system (e.g., UL system), and workmanship.
""",
    "sfrm": """
Write as sprayed fire-resistant material (SFRM) special inspection.

- Focus on substrate condition, thickness, adhesion, and coverage.
""",
    "fire_rated_construction": """
Write as fire-rated construction / partition special inspection.

- Focus on continuity of rated assemblies, joints, and penetrations.
""",
    "structural_stability": """
Write as a structural stability special inspection.

- Focus on temporary bracing, support of existing structure, and impact of work on stability.
""",
    "final": """
Write as a final special inspection summary.

- Summarize status of required special inspections and outstanding items.
""",
    "tenant_protection_plan": """
Write as an inspection of Tenant Protection Plan (TPP) implementation.

- Focus on safeguards for occupied units, egress, dust/noise control, and posted notices.
""",
}


def get_trade_style(trade: str) -> str:
    key = (trade or "").lower()
    return TRADE_STYLE_PROMPTS.get(key, "")


# ============================================================
# Utilities
# ============================================================

def slugify_trade(trade: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", (trade or "").lower()).strip("_")


def slugify_company(company: str) -> str:
    return re.sub(r"[^a-z0-9]+", "_", (company or "").lower()).strip("_") or "default"


def get_drawing_path(drawing_id: str) -> str:
    """
    Resolve a drawing file path by drawing_id.
    """
    for name in os.listdir(UPLOAD_DIR):
        if name.startswith(drawing_id + "."):
            return os.path.join(UPLOAD_DIR, name)
    raise FileNotFoundError(f"No drawing found for id={drawing_id}")


def count_pdf_pages(path: str) -> int:
    """
    Count pages in a PDF using PyPDF2 (if installed) or fallback 1.
    """
    try:
        import PyPDF2
    except ImportError:
        print("PyPDF2 not installed; assuming 1 page for PDF.")
        return 1

    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        return len(reader.pages)


TRADE_CONFIG_PATH = os.path.join(BASE_DIR, "trade_configs.json")


def load_trade_configs() -> Dict[str, Dict[str, str]]:
    if not os.path.exists(TRADE_CONFIG_PATH):
        return {}
    try:
        with open(TRADE_CONFIG_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def save_trade_configs(configs: Dict[str, Dict[str, str]]) -> None:
    with open(TRADE_CONFIG_PATH, "w", encoding="utf-8") as f:
        json.dump(configs, f, indent=2)


def annotation_file_path(drawing_id: str, page: int) -> str:
    return os.path.join(ANNOTATION_DIR, f"{drawing_id}_page{page}.json")


def get_template_for_trade(trade: str, company: Optional[str] = None) -> str:
    """
    Multi-tenant template resolution order:

    1) Company-level master template:
       templates/companies/{company_slug}/Project/report.docx

    2) Global per-trade template:
       templates/global/{trade_slug}_report_template.docx

    3) Default:
       templates/default_report_template.docx
    """
    trade_slug = slugify_trade(trade)
    company_slug = slugify_company(company) if company else None

    # 1) Company-level Project/report.docx
    if company_slug:
        company_project_template = os.path.join(
            COMPANY_TEMPLATE_DIR, company_slug, "Project", "report.docx"
        )
        if os.path.exists(company_project_template):
            return company_project_template

    # 2) Global per-trade
    global_candidate = os.path.join(GLOBAL_TEMPLATE_DIR, f"{trade_slug}_report_template.docx")
    if os.path.exists(global_candidate):
        return global_candidate

    # 3) Default
    if os.path.exists(DEFAULT_TEMPLATE_PATH):
        return DEFAULT_TEMPLATE_PATH

    raise HTTPException(
        status_code=500,
        detail="No DOCX template found. Please upload a template or add default_report_template.docx.",
    )


def replace_placeholders_everywhere(doc: Document, replacements: Dict[str, str]) -> None:
    """
    Replace placeholder tokens in all paragraphs and table cells.
    """

    def _replace_in_paragraph(p):
        if not p.text:
            return
        full_text = "".join(run.text for run in p.runs)
        changed = False
        for old, new in replacements.items():
            if old in full_text:
                full_text = full_text.replace(old, new)
                changed = True
        if changed and p.runs:
            p.runs[0].text = full_text
            for r in p.runs[1:]:
                r.text = ""

    for p in doc.paragraphs:
        _replace_in_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p)


def fill_findings_table(doc: Document, detailed_findings: List[Dict[str, Any]]) -> None:
    """
    Find the first table whose header row includes 'Item' & 'Observation',
    then clear data rows and insert detailed_findings.
    """
    if not doc.tables:
        return

    target_table = None
    for t in doc.tables:
        if not t.rows:
            continue
        hdr_text = " ".join(cell.text for cell in t.rows[0].cells).lower()
        if "item" in hdr_text and "observation" in hdr_text:
            target_table = t
            break

    if not target_table:
        return

    while len(target_table.rows) > 1:
        row = target_table.rows[-1]
        row._element.getparent().remove(row._element)

    for f in detailed_findings or []:
        row_cells = target_table.add_row().cells
        values = [
            f.get("item", ""),
            f.get("observation", ""),
            f.get("code_or_detail_ref", ""),
            f.get("status", ""),
            f.get("remarks", ""),
        ]
        for i, v in enumerate(values):
            if i < len(row_cells):
                row_cells[i].text = v


# ============================================================
# Pydantic models
# ============================================================

class SaveAnnotationsPayload(BaseModel):
    drawing_id: str
    page: int
    annotations: List[Dict[str, Any]]


class GenerateReportPayload(BaseModel):
    drawing_id: str
    trade_type: str
    project_info: str
    page: int
    annotations: List[Dict[str, Any]] = []
    ocr_crop: Optional[Dict[str, Any]] = None
    ocr_text: Optional[str] = ""
    trade_prompt: Optional[str] = ""
    checklist_template: Optional[str] = ""


class SaveTradeConfigPayload(BaseModel):
    trade: str
    system_prompt: str
    checklist_template: str


class ExportPdfPayload(BaseModel):
    pages: List[str]


class ReportDocxPayload(BaseModel):
    report_data: Dict[str, Any]
    company: Optional[str] = None


class TemplateUploadResponse(BaseModel):
    company: str
    trade: str
    path: str


# ============================================================
# Drawing upload
# ============================================================

@app.post("/upload-drawing")
async def upload_drawing(
    file: UploadFile = File(...),
    trade_type: str = Query("welding"),
):
    if not file.filename:
        raise HTTPException(status_code=400, detail="File must have a filename.")

    _, ext = os.path.splitext(file.filename)
    ext = ext.lower() or ".bin"

    drawing_id = str(uuid.uuid4())
    out_name = f"{drawing_id}{ext}"
    out_path = os.path.join(UPLOAD_DIR, out_name)

    contents = await file.read()
    with open(out_path, "wb") as f:
        f.write(contents)

    is_pdf = ext == ".pdf"
    num_pages = 1
    if is_pdf:
        try:
            num_pages = count_pdf_pages(out_path)
        except Exception as e:
            print("Error counting PDF pages:", e)
            num_pages = 1

    return {"drawing_id": drawing_id, "is_pdf": is_pdf, "num_pages": num_pages}


# ============================================================
# Annotations persistence
# ============================================================

@app.post("/save-annotations")
async def save_annotations(payload: SaveAnnotationsPayload):
    path = annotation_file_path(payload.drawing_id, payload.page)
    with open(path, "w", encoding="utf-8") as f:
        json.dump({"annotations": payload.annotations}, f)
    return {"status": "ok"}


@app.get("/load-annotations")
async def load_annotations(
    drawing_id: str = Query(...),
    page: int = Query(1),
):
    path = annotation_file_path(drawing_id, page)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="No annotations for this page")
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data


# ============================================================
# OCR stub
# ============================================================

@app.post("/ocr-crop")
async def ocr_crop(payload: Dict[str, Any] = Body(...)):
    drawing_id = payload.get("drawing_id")
    page = payload.get("page")
    crop = payload.get("crop") or {}
    text = (
        f"[OCR stub] Drawing: {drawing_id}, page {page}, "
        f"crop box: x={crop.get('x')}, y={crop.get('y')}, "
        f"w={crop.get('w')}, h={crop.get('h')}"
    )
    return {"text": text}


# ============================================================
# Trade config
# ============================================================

@app.get("/get-trade-config")
async def get_trade_config(trade: str = Query(...)):
    configs = load_trade_configs()
    cfg = configs.get(trade) or {
        "system_prompt": f"You are a NYC construction special inspector for {trade}.",
        "checklist_template": "• Verify weld/bolt sizes and continuity\n• Verify installation per details\n",
    }
    return cfg


@app.post("/save-trade-config")
async def save_trade_config(payload: SaveTradeConfigPayload):
    configs = load_trade_configs()
    configs[payload.trade] = {
        "system_prompt": payload.system_prompt,
        "checklist_template": payload.checklist_template,
    }
    save_trade_configs(configs)
    return {"status": "ok", "trade": payload.trade}


# ============================================================
# AI report + checklist generation (drawing-based)
# ============================================================

@app.post("/generate-report")
async def generate_report(payload: GenerateReportPayload):
    if not client:
        raise HTTPException(
            status_code=500,
            detail="OpenAI client is not configured. Set OPENAI_API_KEY.",
        )

    base_system = f"""
You are a NYC construction special inspector for {payload.trade_type.upper()}.
Write clear, professional inspection narrative and a concise checklist.

Use:
- Project info
- Drawing page number
- Annotations (marks on drawing)
- OCR text (weld/bolt callouts, gridlines, details)
- Trade-specific style and checklist template

Reference:
- Gridlines and levels (e.g., D/11–12)
- Section/detail numbers (e.g., 5/S3.1)
- Member sizes (e.g., W16x40, PL 3/8")
- Weld/bolt sizes (e.g., 5/16" fillet, 3/4" A325)
"""

    system_prompt = base_system

    style = get_trade_style(payload.trade_type)
    if style:
        system_prompt += "\n\nTRADE-SPECIFIC STYLE:\n" + style

    if payload.trade_prompt:
        system_prompt += "\n\nADDITIONAL TRADE PROMPT:\n" + payload.trade_prompt

    schema_description = """
Return ONLY a single JSON object with this shape:

{
  "report_text": "string",
  "checklist_text": "string",
  "json_items": [
    {
      "item": "string",
      "status": "ACCEPTED | REJECTED | OPEN",
      "note": "string"
    }
  ]
}
"""

    system_prompt += "\n\n" + schema_description + "\nYou MUST respond with strict JSON."

    user_payload = {
        "trade_type": payload.trade_type,
        "project_info": payload.project_info,
        "page": payload.page,
        "annotations": payload.annotations,
        "ocr_text": payload.ocr_text,
        "trade_prompt": payload.trade_prompt,
        "checklist_template": payload.checklist_template,
    }

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)},
    ]

    # ---- OpenAI call with robust error handling ----
    try:
        completion = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=messages,
            temperature=0.2,
        )
        content = completion.choices[0].message.content or ""
    except Exception as e:
        # This is where httpx / network / auth errors will surface
        print("Error calling OpenAI (generate-report):", repr(e))
        raise HTTPException(
            status_code=500,
            detail=f"Error calling OpenAI (generate-report): {e}",
        )

    try:
        data = json.loads(content)
        report_text = data.get("report_text", "")
        checklist_text = data.get("checklist_text", payload.checklist_template or "")
        json_items = data.get("json_items", [])
    except json.JSONDecodeError:
        report_text = content
        checklist_text = payload.checklist_template or ""
        json_items = []

    return {
        "report_text": report_text,
        "checklist_text": checklist_text,
        "json_items": json_items,
    }


# ============================================================
# Export annotated multi-page PDF
# ============================================================

@app.post("/export-pdf")
async def export_pdf(payload: ExportPdfPayload):
    if not payload.pages:
        raise HTTPException(status_code=400, detail="No pages provided.")

    buffer = BytesIO()
    pdf = rl_canvas.Canvas(buffer)

    for data_url in payload.pages:
        m = re.match(r"^data:image/(png|jpeg);base64,(.+)$", data_url)
        if not m:
            continue
        img_bytes = base64.b64decode(m.group(2))
        img = ImageReader(BytesIO(img_bytes))

        width, height = 612, 792
        pdf.setPageSize((width, height))
        pdf.drawImage(img, 0, 0, width=width, height=height)
        pdf.showPage()

    pdf.save()
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="annotated.pdf"'},
    )


# ============================================================
# Generate report from images (vision)
# ============================================================

@app.post("/generate-report-from-images")
async def generate_report_from_images(
    trade: str = Form(...),
    project_name: str = Form(...),
    inspection_date: str = Form(...),
    area_inspected: str = Form(...),
    reference_drawing: str = Form(""),
    reference_detail: str = Form(""),
    reference_details: str = Form(""),
    inspector_notes: str = Form(""),
    main_image: UploadFile = File(...),
    detail_images: Optional[List[UploadFile]] = File(None),
):
    if not client:
        raise HTTPException(
            status_code=500,
            detail="OpenAI client is not configured. Set OPENAI_API_KEY.",
        )

    if not main_image.content_type or not main_image.content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="main_image must be an image file.")

    async def file_to_data_url(f: UploadFile) -> str:
        data = await f.read()
        b64 = base64.b64encode(data).decode("utf-8")
        mime = f.content_type or "image/png"
        return f"data:{mime};base64,{b64}"

    main_img_data_url = await file_to_data_url(main_image)

    detail_data_urls: List[str] = []
    if detail_images:
        for f in detail_images:
            if f.content_type and f.content_type.startswith("image/"):
                detail_data_urls.append(await file_to_data_url(f))

    schema_description = """
Return ONLY a single JSON object with this exact structure:

{
  "project_name": "string",
  "inspection_date": "string",
  "trade": "string",
  "area_inspected": "string",
  "reference_details": "string",
  "overall_summary": "string",
  "detailed_findings": [
    {
      "item": "string",
      "observation": "string",
      "code_or_detail_ref": "string",
      "status": "ACCEPTED | REJECTED | OPEN",
      "remarks": "string"
    }
  ],
  "conclusion": "string"
}
"""

    style = get_trade_style(trade)

    system_prompt = f"""
You are a NYC construction special inspector for the trade: {trade.upper()}.

You are reviewing inspection DRAWING IMAGES (structural steel, welds, bolts, or other trade as applicable).
Using ONLY the provided images + text notes, generate a clear, professional
inspection report in NYC special inspection style.

Focus on:
- Member / system details (e.g., W16x40, PL 3/8", duct sizes, pipe sizes, etc.)
- Weld / bolt sizes, or trade-specific inspection items
- Gridlines and locations (e.g., grid D/11–12, level references)
- Whether work is ACCEPTED, REJECTED, or OPEN and why.

{style}

{schema_description}
You MUST respond with strict JSON only, no extra commentary.
"""

    text_block = f"""
Project name: {project_name}
Inspection date: {inspection_date}
Trade: {trade}
Area inspected: {area_inspected}
Reference drawing: {reference_drawing}
Reference detail: {reference_detail}
Additional references: {reference_details}
Inspector notes: {inspector_notes}

The following images show the area inspected and any related details.
Use them to understand member types, weld/bolt callouts, MEP components,
or other trade-specific conditions as visible.
"""

    user_content: List[Dict[str, Any]] = [
        {"type": "text", "text": text_block},
        {"type": "image_url", "image_url": {"url": main_img_data_url, "detail": "high"}},
    ]

    for url in detail_data_urls:
        user_content.append(
            {"type": "image_url", "image_url": {"url": url, "detail": "high"}}
        )

    # ---- OpenAI vision call with robust error handling ----
    try:
        completion = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
            temperature=0.2,
        )
        content = completion.choices[0].message.content or ""
    except Exception as e:
        print("Error calling OpenAI (images):", repr(e))
        raise HTTPException(
            status_code=500,
            detail=f"Error calling OpenAI (images): {e}",
        )

    try:
        report_data = json.loads(content)
    except json.JSONDecodeError:
        report_data = {
            "project_name": project_name,
            "inspection_date": inspection_date,
            "trade": trade,
            "area_inspected": area_inspected,
            "reference_details": reference_details,
            "overall_summary": content,
            "detailed_findings": [],
            "conclusion": "",
        }

    # Ensure basic fields present
    report_data.setdefault("project_name", project_name)
    report_data.setdefault("inspection_date", inspection_date)
    report_data.setdefault("trade", trade)
    report_data.setdefault("area_inspected", area_inspected)
    report_data.setdefault("reference_details", reference_details)

    report_data["reference_drawing"] = reference_drawing
    report_data["reference_detail"] = reference_detail

    return JSONResponse({"report_data": report_data})


# ============================================================
# Upload DOCX template for company + trade (optional)
# ============================================================

@app.post("/upload-report-template", response_model=TemplateUploadResponse)
async def upload_report_template(
    company: str = Form(...),
    trade: str = Form(...),
    file: UploadFile = File(...),
):
    if not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx templates are supported.")

    company_slug = slugify_company(company)
    trade_slug = slugify_trade(trade)

    company_dir = os.path.join(COMPANY_TEMPLATE_DIR, company_slug)
    os.makedirs(company_dir, exist_ok=True)

    dest_path = os.path.join(company_dir, f"{trade_slug}_report_template.docx")

    contents = await file.read()
    with open(dest_path, "wb") as f:
        f.write(contents)

    return TemplateUploadResponse(company=company, trade=trade, path=dest_path)


# ============================================================
# Export DOCX from JSON report
# ============================================================

@app.post("/export-report-docx")
async def export_report_docx(payload: ReportDocxPayload):
    """
    Map JSON report_data into a DOCX template (multi-tenant).
    """
    rd = payload.report_data or {}

    company = payload.company or rd.get("company")
    trade = rd.get("trade", "general")

    template_path = get_template_for_trade(trade, company=company)

    try:
        doc = Document(template_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error loading DOCX template: {e}")

    detailed_findings = rd.get("detailed_findings", []) or []

    # Auto-derive deficiencies
    defs = rd.get("deficiencies")
    if defs is None:
        defs = []
        for df in detailed_findings:
            status = (df.get("status") or "").upper()
            if status and status != "ACCEPTED":
                defs.append(
                    {
                        "no": str(len(defs) + 1),
                        "text": df.get("observation", "") or df.get("item", ""),
                    }
                )
        rd["deficiencies"] = defs

    if not rd.get("deficiencies_summary"):
        if defs:
            rd["deficiencies_summary"] = "Deficiencies noted"
        else:
            rd["deficiencies_summary"] = "No deficiencies noted"

    # Auto-derive observations
    obs_list = rd.get("observations")
    if obs_list is None or len(obs_list) == 0:
        obs_list = []
        for df in detailed_findings:
            obs_list.append(
                {
                    "general_location": rd.get("area_inspected", ""),
                    "specific_location": df.get("item", ""),
                    "system_or_element": df.get("observation", ""),
                    "su": (df.get("status") or "").upper(),
                    "remarks": df.get("remarks", ""),
                }
            )
        rd["observations"] = obs_list
    else:
        obs_list = obs_list or []

    replacements: Dict[str, str] = {
        "{{PROJECT_NAME}}": rd.get("project_name", ""),
        "{{INSPECTION_DATE}}": rd.get("inspection_date", ""),
        "{{TRADE}}": rd.get("trade", ""),
        "{{AREA_INSPECTED}}": rd.get("area_inspected", ""),
        "{{REFERENCE_DETAILS}}": rd.get("reference_details", ""),
        "{{OVERALL_SUMMARY}}": rd.get("overall_summary", ""),
        "{{CONCLUSION}}": rd.get("conclusion", ""),
    }

    replacements.update({
        "{{PROJECT_NUMBER}}": rd.get("project_number", ""),
        "{{CLIENT_NAME}}": rd.get("client_name", ""),
        "{{BIS_NUMBER}}": rd.get("bis_number", ""),
        "{{GC_CM}}": rd.get("gc_cm", ""),
        "{{ARCHITECT}}": rd.get("architect", ""),
        "{{ENGINEER}}": rd.get("engineer", ""),
        "{{TIME_IN}}": rd.get("time_in", ""),
        "{{TIME_OUT}}": rd.get("time_out", ""),
        "{{REPORT_NUMBER}}": rd.get("report_number", ""),
        "{{INSPECTORS}}": rd.get("inspectors", ""),
        "{{REPORTED_TO}}": rd.get("reported_to", ""),
        "{{INSPECTIONS_LIST}}": rd.get("inspections_list", ""),
        "{{PERSONS_PRESENT}}": rd.get("persons_present", ""),
        "{{SITE_REMARKS}}": rd.get("site_remarks", ""),
        "{{DEFICIENCIES_SUMMARY}}": rd.get("deficiencies_summary", ""),
        "{{REINSPECTION_REQUIRED}}": rd.get("reinspection_required", ""),
        "{{TOTAL_ATTACHMENTS}}": rd.get("total_attachments", ""),
        "{{ATTACHMENTS_LIST}}": rd.get("attachments_list", ""),
        "{{OTHER_NOTES}}": rd.get("other_notes", ""),
        "{{REFERENCE_DRAWING}}": rd.get("reference_drawing", ""),
        "{{REFERENCE_DETAIL}}": rd.get("reference_detail", ""),
    })

    # Photos placeholders
    photos = rd.get("photos", []) or []
    for i in range(1, 4):
        p = photos[i - 1] if i - 1 < len(photos) else {}
        replacements[f"{{{{PHOTO_{i}_TITLE}}}}"] = p.get("title", "")
        replacements[f"{{{{PHOTO_{i}_NOTE}}}}"] = p.get("note", "")

    # Deficiencies placeholders
    defs = rd.get("deficiencies", []) or []
    for i in range(1, 5 + 1):
        d = defs[i - 1] if i - 1 < len(defs) else {}
        replacements[f"{{{{DEF_{i}_NO}}}}"] = d.get("no", "")
        replacements[f"{{{{DEF_{i}_TEXT}}}}"] = d.get("text", "")

    # Previous deficiencies resolved
    prev_defs = rd.get("previous_deficiencies_resolved", []) or []
    for i in range(1, 3 + 1):
        d = prev_defs[i - 1] if i - 1 < len(prev_defs) else {}
        replacements[f"{{{{PREV_DEF_{i}_NO}}}}"] = d.get("no", "")
        replacements[f"{{{{PREV_DEF_{i}_RESOLUTION}}}}"] = d.get("resolution", "")

    # Observations table placeholders
    obs_list = obs_list or []
    for i in range(1, 15 + 1):
        o = obs_list[i - 1] if i - 1 < len(obs_list) else {}
        replacements[f"{{{{OBS_{i}_GENERAL}}}}"] = o.get("general_location", "")
        replacements[f"{{{{OBS_{i}_SPECIFIC}}}}"] = o.get("specific_location", "")
        replacements[f"{{{{OBS_{i}_SYSTEM}}}}"] = o.get("system_or_element", "")
        replacements[f"{{{{OBS_{i}_SU}}}}"] = o.get("su", "")
        replacements[f"{{{{OBS_{i}_REMARKS}}}}"] = o.get("remarks", "")

    replace_placeholders_everywhere(doc, replacements)

    if detailed_findings:
        try:
            fill_findings_table(doc, detailed_findings)
        except Exception as e:
            print("fill_findings_table error:", e)

    # Embed annotated images
    attachments_images = rd.get("attachments_images", []) or []
    if attachments_images:
        doc.add_page_break()
        doc.add_paragraph("Photo Attachments", style=None)
        for idx, data_url in enumerate(attachments_images, start=1):
            try:
                m = re.match(r"^data:image/(png|jpeg);base64,(.+)$", data_url)
                if not m:
                    continue
                img_bytes = base64.b64decode(m.group(2))
                img_stream = BytesIO(img_bytes)
                doc.add_paragraph(f"Photo {idx}:")
                doc.add_picture(img_stream, width=Inches(5.5))
            except Exception as e:
                print("Error embedding attachment image:", e)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    safe_name = slugify_trade(rd.get("project_name") or "inspection_report")
    filename = f"{safe_name}.docx"

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ============================================================
# Image-based report UI (HTML+JS)
# ============================================================
@app.get("/image-report-ui", response_class=HTMLResponse)
async def image_report_ui():
    html = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8" />
    <title>Image-Based Special Inspection Report</title>
    <meta name="viewport" content="width=device-width, initial-scale=1" />
    <style>
        :root {
            font-family: system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
            color-scheme: light dark;
        }
        body {
            margin: 0;
            padding: 0;
            background: #0f172a;
            color: #e5e7eb;
        }
        .page {
            max-width: 1100px;
            margin: 0 auto;
            padding: 24px 16px 48px;
        }
        h1 {
            margin-top: 0;
            font-size: 1.8rem;
        }
        h2 {
            font-size: 1.3rem;
            margin-top: 1.5rem;
        }
        .card {
            background: #020617;
            border-radius: 12px;
            padding: 16px 20px;
            border: 1px solid #1e293b;
            box-shadow: 0 10px 35px rgba(15, 23, 42, 0.8);
        }
        .grid {
            display: grid;
            grid-template-columns: 1.2fr 1fr;
            gap: 18px;
        }
        .grid-2 {
            display: grid;
            grid-template-columns: repeat(2, minmax(0, 1fr));
            gap: 12px;
        }
        label {
            display: block;
            font-size: 0.85rem;
            margin-bottom: 4px;
            color: #cbd5f5;
        }
        input[type="text"],
        input[type="date"],
        select,
        textarea {
            width: 100%;
            padding: 8px 10px;
            border-radius: 8px;
            border: 1px solid #1f2937;
            background: #020617;
            color: #e5e7eb;
            font-size: 0.9rem;
            box-sizing: border-box;
        }
        textarea {
            resize: vertical;
            min-height: 80px;
        }
        input[type="file"] {
            width: 100%;
            border-radius: 8px;
            border: 1px dashed #334155;
            padding: 10px;
            background: #020617;
            color: #e5e7eb;
            font-size: 0.9rem;
            box-sizing: border-box;
        }
        .field {
            margin-bottom: 10px;
        }
        .btn-row {
            margin-top: 8px;
            display: flex;
            align-items: center;
            gap: 10px;
        }
        button {
            border-radius: 999px;
            border: none;
            padding: 8px 18px;
            background: #2563eb;
            color: #f9fafb;
            font-size: 0.9rem;
            font-weight: 500;
            cursor: pointer;
            display: inline-flex;
            align-items: center;
            gap: 6px;
        }
        button:disabled {
            opacity: 0.6;
            cursor: default;
        }
        .pill {
            display: inline-flex;
            align-items: center;
            gap: 6px;
            padding: 3px 8px;
            border-radius: 999px;
            border: 1px solid #1f2937;
            font-size: 0.7rem;
            color: #9ca3af;
        }
        .status-dot {
            width: 8px;
            height: 8px;
            border-radius: 999px;
        }
        .status-ACCEPTED { background: #22c55e; }
        .status-REJECTED { background: #ef4444; }
        .status-OPEN { background: #eab308; }
        .small {
            font-size: 0.8rem;
            color: #9ca3af;
        }
        .report-block {
            border-radius: 10px;
            border: 1px solid #1f2937;
            padding: 10px 12px;
            background: #020617;
            font-size: 0.9rem;
            white-space: pre-wrap;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            font-size: 0.85rem;
            margin-top: 4px;
        }
        th, td {
            border: 1px solid #1f2937;
            padding: 6px 8px;
            vertical-align: top;
        }
        th {
            background: #020617;
            font-weight: 500;
            text-align: left;
        }
        .badge {
            display: inline-flex;
            align-items: center;
            gap: 4px;
            border-radius: 999px;
            padding: 2px 8px;
            font-size: 0.7rem;
            border: 1px solid #4b5563;
        }
        .fade {
            opacity: 0.6;
        }
        .error {
            color: #fecaca;
            font-size: 0.85rem;
            margin-top: 4px;
        }
        @media (max-width: 900px) {
            .grid {
                grid-template-columns: 1fr;
            }
        }
    </style>
</head>
<body>
<div class="page">
    <h1>Image-Based Special Inspection Report</h1>
    <p class="small fade">
        Upload one <strong>main drawing/connection image</strong> plus optional
        <strong>detail images</strong>. The backend will call
        <code>/generate-report-from-images</code> and use GPT-Vision to draft a
        professional special inspection report, then you can export it to DOCX.
    </p>

    <div class="card">
        <div class="grid">
            <!-- LEFT: FORM -->
            <div>
                <h2>1. Inspection Inputs</h2>
                <form id="image-report-form" enctype="multipart/form-data">
                    <div class="grid-2">
                        <div class="field">
                            <label for="company_name">Company / Template</label>
                            <select id="company_name" name="company_name" required>
                                <option value="default">Default Template</option>
                                <option value="NYC Structural Reports">NYC Structural Reports</option>
                                <option value="Acme Steel Inspections">Acme Steel Inspections</option>
                            </select>
                            <div class="small fade">
                                Select which company's Project/report.docx template to use.
                            </div>
                        </div>
                        <div class="field">
                            <label for="trade">Trade</label>
                            <select id="trade" name="trade" required>
                                <option value="welding">Welding</option>
                                <option value="bolting">Bolting / HSB</option>
                                <option value="welding_bolting">Welding + Bolting</option>
                                <option value="detail">Detail</option>
                                <option value="cold_formed_steel">Cold formed steel</option>
                                <option value="sprinkler">Sprinkler</option>
                                <option value="standpipe">Standpipe</option>
                                <option value="hvac">HVAC</option>
                                <option value="fire_resistance_penetration">Fire resistance penetration</option>
                                <option value="sfrm">Sprayed fire-resistant material</option>
                                <option value="fire_rated_construction">Fire rated construction</option>
                                <option value="structural_stability">Structural stability</option>
                                <option value="final">Final</option>
                                <option value="tenant_protection_plan">Tenant protection plan</option>
                            </select>
                        </div>
                    </div>

                    <div class="grid-2">
                        <div class="field">
                            <label for="project_name">Project name</label>
                            <input type="text" id="project_name" name="project_name"
                                   placeholder="114 5th Ave FL#10" required />
                        </div>
                        <div class="field">
                            <label for="inspection_date">Inspection date</label>
                            <input type="date" id="inspection_date" name="inspection_date" required />
                        </div>
                    </div>

                    <div class="field">
                        <label for="area_inspected">Area inspected (gridlines / level)</label>
                        <input type="text" id="area_inspected" name="area_inspected"
                               placeholder="10TH FLOOR – Grid D/11–12" required />
                    </div>

                    <div class="grid-2">
                        <div class="field">
                            <label for="reference_drawing">Reference Drawing #</label>
                            <input type="text" id="reference_drawing" name="reference_drawing"
                                   placeholder="D1/E1-1" />
                        </div>
                        <div class="field">
                            <label for="reference_detail">Reference Detail #</label>
                            <input type="text" id="reference_detail" name="reference_detail"
                                   placeholder="Detail 11-ac1" />
                        </div>
                    </div>

                    <div class="field">
                        <label for="reference_details">Additional reference notes</label>
                        <input type="text" id="reference_details" name="reference_details"
                               placeholder="Other details / sections / notes" />
                    </div>

                    <div class="field">
                        <label for="inspector_notes">On-site inspector notes</label>
                        <textarea id="inspector_notes" name="inspector_notes"
                                  placeholder="Welds appear continuous; no visible undercut; bolts snug-tight; no rotation of plies."></textarea>
                    </div>

                    <hr style="border: 0; border-top: 1px solid #1f2937; margin: 12px 0;" />

                    <div class="field">
                        <label for="main_image">
                            Main image (required)
                            <span class="small fade">– drawing crop / connection overview</span>
                        </label>
                        <input type="file" id="main_image" name="main_image" accept="image/*" required />
                    </div>

                    <div class="field">
                        <label for="detail_images">
                            Detail images (optional)
                            <span class="small fade">– callouts, close-ups, detail bubbles</span>
                        </label>
                        <input type="file" id="detail_images" name="detail_images"
                               accept="image/*" multiple />
                    </div>

                    <hr style="border: 0; border-top: 1px solid #1f2937; margin: 12px 0;" />

                    <div class="field">
                        <label>
                            Annotate main image
                            <span class="small fade">– mark inspected beams, welds, bolts, and add text labels</span>
                        </label>
                        <div class="small fade" style="margin-bottom:6px;">
                            When you select a <strong>Main image</strong> above, it will appear on this canvas.
                            Use the tools to mark inspected beams, welds, bolts, or details. This annotated
                            image will be attached to the DOCX report.
                        </div>

                        <div style="display:flex; flex-wrap:wrap; gap:8px; align-items:center; margin-bottom:6px;">
                            <label class="small">
                                Tool:
                                <select id="draw_tool" style="margin-left:4px;">
                                    <option value="none">View only</option>
                                    <option value="pen">Freehand pen</option>
                                    <option value="rect">Rectangle</option>
                                    <option value="text">Text label</option>
                                </select>
                            </label>

                            <label class="small">
                                Color:
                                <input type="color" id="draw_color" value="#ef4444" />
                            </label>

                            <label class="small">
                                Width:
                                <input type="number" id="draw_width" value="3" min="1" max="12" style="width:50px;" />
                            </label>

                            <button type="button" id="clear_canvas_btn">Clear</button>

                            <a href="http://localhost:3000" target="_blank" class="small fade"
                               style="margin-left:auto; text-decoration:none;">
                                🔗 Open full PDF editor
                            </a>
                        </div>

                        <canvas id="annot_main_canvas"
                                style="width:100%; max-width:520px; border:1px solid #1f2937; border-radius:8px; background:#020617;"></canvas>
                    </div>

                    <div class="btn-row">
                        <button type="submit" id="submit-btn">
                            <span id="btn-label">⚙️ Generate AI Report from Images</span>
                        </button>
                        <button type="button" id="export-docx-btn" class="secondary">
                            ⬇️ Export DOCX
                        </button>
                        <span id="status-pill" class="pill">
                            <span class="status-dot status-OPEN"></span>
                            <span id="status-text">Idle</span>
                        </span>
                    </div>
                    <div id="form-error" class="error" style="display:none;"></div>
                </form>
            </div>

            <!-- RIGHT: REPORT PREVIEW -->
            <div>
                <h2>2. Report Preview</h2>
                <p class="small fade">
                    After generation, the AI output appears here. You can then click
                    <strong>Export DOCX</strong> to download a report using the selected
                    company template.
                </p>

                <div class="field">
                    <label>Overall summary</label>
                    <div id="overall_summary" class="report-block small fade">
                        No report generated yet.
                    </div>
                </div>

                <div class="field">
                    <label>Detailed findings</label>
                    <div id="findings_container" class="report-block">
                        <span class="small fade">No findings yet.</span>
                    </div>
                </div>

                <div class="field">
                    <label>Conclusion</label>
                    <div id="conclusion" class="report-block small fade">
                        No conclusion yet.
                    </div>
                </div>

                <div class="field">
                    <label class="small">Raw JSON (for debugging / DOCX mapping)</label>
                    <textarea id="raw_json" readonly class="small"
                              style="min-height: 80px;"></textarea>
                </div>
            </div>
        </div>
    </div>
</div>

<script>
(function() {
    const form = document.getElementById("image-report-form");
    const submitBtn = document.getElementById("submit-btn");
    const exportDocxBtn = document.getElementById("export-docx-btn");
    const btnLabel = document.getElementById("btn-label");
    const statusText = document.getElementById("status-text");
    const formError = document.getElementById("form-error");
    const companySelect = document.getElementById("company_name");

    const overallSummaryEl = document.getElementById("overall_summary");
    const findingsContainer = document.getElementById("findings_container");
    const conclusionEl = document.getElementById("conclusion");
    const rawJsonEl = document.getElementById("raw_json");

    let lastReportData = null;

    // --- Canvas annotation state ---
    const mainImageInput = document.getElementById("main_image");
    const detailImagesInput = document.getElementById("detail_images");
    const canvas = document.getElementById("annot_main_canvas");
    const drawTool = document.getElementById("draw_tool");
    const drawColor = document.getElementById("draw_color");
    const drawWidth = document.getElementById("draw_width");
    const clearCanvasBtn = document.getElementById("clear_canvas_btn");
    const ctx = canvas.getContext("2d");

    let baseImage = null;
    let isDrawing = false;
    let startX = 0;
    let startY = 0;
    let lastX = 0;
    let lastY = 0;

    // Store detail images as data URLs so we can send them as attachments
    let detailImageDataURLs = [];

    function setStatus(text) {
        statusText.textContent = text;
    }

    function setLoading(isLoading) {
        submitBtn.disabled = isLoading;
        if (isLoading) {
            btnLabel.textContent = "⏳ Generating...";
            setStatus("Calling /generate-report-from-images");
        } else {
            btnLabel.textContent = "⚙️ Generate AI Report from Images";
        }
    }

    function redrawBase() {
        if (!baseImage) return;
        ctx.clearRect(0, 0, canvas.width, canvas.height);
        ctx.drawImage(baseImage, 0, 0, canvas.width, canvas.height);
    }

    function loadMainImageToCanvas(file) {
        const reader = new FileReader();
        reader.onload = function(e) {
            const img = new Image();
            img.onload = function() {
                baseImage = img;
                const maxWidth = 800;
                const scale = Math.min(1, maxWidth / img.width);
                const w = img.width * scale;
                const h = img.height * scale;
                canvas.width = w;
                canvas.height = h;
                redrawBase();
            };
            img.src = e.target.result;
        };
        reader.readAsDataURL(file);
    }

    if (mainImageInput) {
        mainImageInput.addEventListener("change", function() {
            const file = mainImageInput.files && mainImageInput.files[0];
            if (file) {
                loadMainImageToCanvas(file);
            }
        });
    }

    // Read detail images as data URLs for later attachment
    if (detailImagesInput) {
        detailImagesInput.addEventListener("change", function() {
            detailImageDataURLs = [];
            const files = Array.from(detailImagesInput.files || []);
            files.forEach(file => {
                if (file.type && file.type.startsWith("image/")) {
                    const reader = new FileReader();
                    reader.onload = (e) => {
                        detailImageDataURLs.push(e.target.result);
                    };
                    reader.readAsDataURL(file);
                }
            });
        });
    }

    function getCanvasPos(evt) {
        const rect = canvas.getBoundingClientRect();
        return {
            x: (evt.clientX - rect.left) * (canvas.width / rect.width),
            y: (evt.clientY - rect.top) * (canvas.height / rect.height),
        };
    }

    canvas.addEventListener("mousedown", function (e) {
        if (!baseImage) return;
        const tool = drawTool.value;

        if (tool === "none") return;

        const pos = getCanvasPos(e);

        if (tool === "text") {
            // --- TEXT TOOL: click to place text ---
            const t = window.prompt("Enter annotation text:");
            if (t && t.trim().length > 0) {
                ctx.save();
                ctx.fillStyle = drawColor.value || "#ef4444";
                ctx.font = "18px system-ui"; // Bigger, more readable text
                ctx.shadowColor = "rgba(0,0,0,0.7)";
                ctx.shadowBlur = 3;
                ctx.fillText(t, pos.x, pos.y);
                ctx.restore();
            }
            return;
        }

        // pen or rect: start drawing
        isDrawing = true;
        startX = pos.x;
        startY = pos.y;
        lastX = pos.x;
        lastY = pos.y;

        if (tool === "pen") {
            ctx.strokeStyle = drawColor.value || "#ef4444";
            ctx.lineWidth = parseFloat(drawWidth.value) || 3;
            ctx.lineCap = "round";
        }
    });

    canvas.addEventListener("mousemove", function (e) {
        if (!isDrawing || !baseImage) return;
        const tool = drawTool.value;
        const pos = getCanvasPos(e);

        if (tool === "pen") {
            ctx.strokeStyle = drawColor.value || "#ef4444";
            ctx.lineWidth = parseFloat(drawWidth.value) || 3;
            ctx.lineCap = "round";
            ctx.beginPath();
            ctx.moveTo(lastX, lastY);
            ctx.lineTo(pos.x, pos.y);
            ctx.stroke();
            lastX = pos.x;
            lastY = pos.y;
        } else if (tool === "rect") {
            // preview rectangle
            redrawBase();
            ctx.strokeStyle = drawColor.value || "#ef4444";
            ctx.lineWidth = parseFloat(drawWidth.value) || 3;
            ctx.setLineDash([6, 4]);
            ctx.strokeRect(startX, startY, pos.x - startX, pos.y - startY);
            ctx.setLineDash([]);
            lastX = pos.x;
            lastY = pos.y;
        }
    });

    window.addEventListener("mouseup", function () {
        if (!isDrawing) return;
        const tool = drawTool.value;

        if (tool === "rect" && baseImage) {
            // final rectangle (solid) using lastX/lastY
            redrawBase();
            ctx.strokeStyle = drawColor.value || "#ef4444";
            ctx.lineWidth = parseFloat(drawWidth.value) || 3;
            ctx.setLineDash([]);
            ctx.strokeRect(
                startX,
                startY,
                lastX - startX,
                lastY - startY
            );
        }
        isDrawing = false;
    });

    if (clearCanvasBtn) {
        clearCanvasBtn.addEventListener("click", function () {
            if (!baseImage) {
                ctx.clearRect(0, 0, canvas.width, canvas.height);
                return;
            }
            redrawBase();
        });
    }

    function renderFindings(detailedFindings) {
        if (!Array.isArray(detailedFindings) || detailedFindings.length === 0) {
            findingsContainer.innerHTML =
                '<span class="small fade">No detailed findings in response.</span>';
            return;
        }

        let html = '<table><thead><tr>' +
            '<th>Item</th><th>Observation</th><th>Ref</th><th>Status</th><th>Remarks</th>' +
            '</tr></thead><tbody>';

        for (const f of detailedFindings) {
            const status = (f.status || f.Status || "").toString().toUpperCase();
            const cls = status ? "status-" + status : "";
            html += "<tr>";
            html += "<td>" + (f.item || f.Item || "") + "</td>";
            html += "<td>" + (f.observation || f.Observation || "") + "</td>";
            html += "<td>" + (f.code_or_detail_ref || f.Ref || "") + "</td>";
            html += '<td><span class="badge">' +
                '<span class="status-dot ' + cls + '"></span>' +
                (status || "N/A") +
                "</span></td>";
            html += "<td>" + (f.remarks || f.Remarks || "") + "</td>";
            html += "</tr>";
        }

        html += "</tbody></table>";
        findingsContainer.innerHTML = html;
    }

    form.addEventListener("submit", async function (e) {
        e.preventDefault();
        formError.style.display = "none";
        formError.textContent = "";
        setLoading(true);

        const formData = new FormData(form);

        try {
            const resp = await fetch("/generate-report-from-images", {
                method: "POST",
                body: formData
            });

            if (!resp.ok) {
                const text = await resp.text();
                throw new Error("HTTP " + resp.status + ": " + text);
            }

            const data = await resp.json();
            const report = data.report_data || data;

            lastReportData = report;

            overallSummaryEl.classList.remove("fade");
            conclusionEl.classList.remove("fade");

            overallSummaryEl.textContent = report.overall_summary || "(no overall_summary)";
            conclusionEl.textContent = report.conclusion || "(no conclusion)";

            renderFindings(report.detailed_findings);

            rawJsonEl.value = JSON.stringify(report, null, 2);

            setStatus("Report generated");
        } catch (err) {
            console.error(err);
            formError.style.display = "block";
            formError.textContent = "Error generating report: " + err.message;
            setStatus("Error");
        } finally {
            setLoading(false);
        }
    });

    exportDocxBtn.addEventListener("click", async function () {
        if (!lastReportData) {
            alert("Generate a report first.");
            return;
        }

        // Capture annotated main image as attachment (if canvas has content)
        let attachmentsImages = [];
        if (canvas && canvas.width > 0 && canvas.height > 0) {
            try {
                const dataUrl = canvas.toDataURL("image/png");
                attachmentsImages.push(dataUrl);
            } catch (e) {
                console.warn("Could not capture canvas image:", e);
            }
        }

        // Also attach detail images
        if (detailImageDataURLs.length > 0) {
            // You can limit if needed, e.g., slice(0, 6)
            attachmentsImages = attachmentsImages.concat(detailImageDataURLs);
        }

        const payloadReport = Object.assign({}, lastReportData);
        if (attachmentsImages.length > 0) {
            payloadReport.attachments_images = attachmentsImages;
            if (!payloadReport.total_attachments) {
                payloadReport.total_attachments = String(attachmentsImages.length);
            }
            if (!payloadReport.attachments_list) {
                payloadReport.attachments_list = "Annotated photo(s) and detail images attached.";
            }
        }

        const companyName = companySelect ? companySelect.value : "default";
        payloadReport.company = companyName;

        try {
            const resp = await fetch("/export-report-docx", {
                method: "POST",
                headers: {"Content-Type": "application/json"},
                body: JSON.stringify({
                    report_data: payloadReport,
                    company: companyName
                })
            });
            if (!resp.ok) {
                const text = await resp.text();
                throw new Error("HTTP " + resp.status + ": " + text);
            }
            const blob = await resp.blob();
            const url = URL.createObjectURL(blob);
            const a = document.createElement("a");
            a.href = url;
            a.download = "inspection_report.docx";
            document.body.appendChild(a);
            a.click();
            a.remove();
            URL.revokeObjectURL(url);
        } catch (err) {
            console.error(err);
            alert("Error exporting DOCX: " + err.message);
        }
    });
})();
</script>

</body>
</html>
    """
    return HTMLResponse(html)
